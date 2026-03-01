#!/usr/bin/env python3
"""
Tier B Parser — Semi-structured Data (Word/PDF)

Parses document files (.docx, .pdf) with text extraction and field mapping:
  - 심방일지.docx  -> member visit history entries
  - 회의록.docx    -> meeting minutes (governance records)
  - General .docx/.pdf -> text extraction with field detection

Uses church-glossary.yaml for Korean term normalization.
Produces JSON staging files in inbox/staging/.

Usage:
    python3 tier_b_parser.py --file <path> --glossary <glossary.yaml> --staging-dir <dir>
"""

import argparse
import json
import os
import re
import sys
from datetime import datetime

import yaml

# ---------------------------------------------------------------------------
# Optional dependency: python-docx (graceful fallback)
# ---------------------------------------------------------------------------
try:
    import docx
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------
DATE_RE = re.compile(r"\d{4}[-./]\d{1,2}[-./]\d{1,2}")
DATE_KO_RE = re.compile(r"(\d{4})년\s*(\d{1,2})월\s*(\d{1,2})일")
PHONE_RE = re.compile(r"010[-.\s]?\d{4}[-.\s]?\d{4}")
MEMBER_ID_RE = re.compile(r"M\d{3,}")

# Known file-to-schema mappings
FILE_ROUTING = {
    "심방일지": {
        "target": "data/members.yaml",
        "section": "history",
        "parser": "parse_visitation",
    },
    "회의록": {
        "target": "data/schedule.yaml",
        "section": "meeting_minutes",
        "parser": "parse_meeting_minutes",
    },
}

# Visitation field patterns (Korean)
VISIT_FIELD_PATTERNS = {
    "방문자": re.compile(r"방문자\s*[:：]\s*(.+)"),
    "방문대상": re.compile(r"(?:방문\s*대상|심방\s*대상|교인)\s*[:：]\s*(.+)"),
    "일시": re.compile(r"(?:일시|날짜|방문일)\s*[:：]\s*(.+)"),
    "장소": re.compile(r"장소\s*[:：]\s*(.+)"),
    "사유": re.compile(r"(?:사유|목적|방문\s*사유)\s*[:：]\s*(.+)"),
    "내용": re.compile(r"(?:내용|기록|심방\s*내용)\s*[:：]\s*(.+)"),
    "기도제목": re.compile(r"(?:기도\s*제목|기도)\s*[:：]\s*(.+)"),
    "후속조치": re.compile(r"(?:후속\s*조치|후속|다음\s*단계)\s*[:：]\s*(.+)"),
}


# ---------------------------------------------------------------------------
# Glossary loader
# ---------------------------------------------------------------------------
def load_glossary(glossary_path):
    """Load church-glossary.yaml and build Korean->English lookup."""
    if not os.path.isfile(glossary_path):
        return {}
    with open(glossary_path, "r", encoding="utf-8") as f:
        data = yaml.safe_load(f) or {}
    lookup = {}
    for term in data.get("terms", []):
        korean = term.get("korean", "")
        english = term.get("english", "")
        if korean and english:
            lookup[korean] = english
    return lookup


# ---------------------------------------------------------------------------
# Text extraction
# ---------------------------------------------------------------------------
def extract_text_docx(filepath):
    """Extract text from a .docx file.

    Returns list of paragraphs as strings.
    Raises RuntimeError if python-docx is not installed.
    """
    if not HAS_DOCX:
        raise RuntimeError(
            "python-docx is not installed. Install it with: pip install python-docx\n"
            "This is required to parse .docx files in Tier B.\n"
            "Alternatively, save the document as plain text (.txt) and "
            "place it in the inbox."
        )
    doc = docx.Document(filepath)
    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    # Also extract from tables
    for table in doc.tables:
        for row in table.rows:
            row_texts = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_texts.append(cell_text)
            if row_texts:
                paragraphs.append(" | ".join(row_texts))

    return paragraphs


def extract_text_pdf(filepath):
    """Extract text from a PDF file.

    Uses basic binary text extraction as a fallback when no PDF library is
    available. For production use, install PyPDF2 or pdfplumber.

    Returns list of text chunks.
    """
    # Try PyPDF2 first
    try:
        import PyPDF2
        chunks = []
        with open(filepath, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    for line in text.split("\n"):
                        line = line.strip()
                        if line:
                            chunks.append(line)
        return chunks
    except ImportError:
        pass

    # Try pdfplumber
    try:
        import pdfplumber
        chunks = []
        with pdfplumber.open(filepath) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    for line in text.split("\n"):
                        line = line.strip()
                        if line:
                            chunks.append(line)
        return chunks
    except ImportError:
        pass

    # Fallback: basic binary text extraction (limited accuracy)
    chunks = []
    try:
        with open(filepath, "rb") as f:
            raw = f.read()
        # Try to decode UTF-8 text segments
        text = raw.decode("utf-8", errors="ignore")
        for line in text.split("\n"):
            line = line.strip()
            # Filter out binary noise: require at least some Korean or ASCII text
            if line and (re.search(r"[\uac00-\ud7a3]", line) or re.search(r"[a-zA-Z]{3,}", line)):
                chunks.append(line)
    except Exception:
        pass

    if not chunks:
        raise RuntimeError(
            f"Cannot extract text from PDF '{filepath}'. "
            "Install one of: pip install PyPDF2  OR  pip install pdfplumber"
        )
    return chunks


# ---------------------------------------------------------------------------
# Date normalization
# ---------------------------------------------------------------------------
def normalize_date(text):
    """Extract and normalize date from text.

    Returns (YYYY-MM-DD string, confidence) or (None, 0.0).
    """
    if not text:
        return None, 0.0
    text = str(text).strip()

    # Korean date pattern: 2026년 2월 28일
    m = DATE_KO_RE.search(text)
    if m:
        return f"{m.group(1)}-{int(m.group(2)):02d}-{int(m.group(3)):02d}", 0.95

    # ISO-like: 2026-02-28 or 2026/02/28 or 2026.02.28
    m = DATE_RE.search(text)
    if m:
        parts = re.split(r"[-./]", m.group(0))
        if len(parts) == 3:
            return f"{parts[0]}-{int(parts[1]):02d}-{int(parts[2]):02d}", 0.9

    return None, 0.0


# ---------------------------------------------------------------------------
# Schema-specific parsers
# ---------------------------------------------------------------------------
def parse_visitation(paragraphs, glossary, source_file):
    """Parse visitation report (심방일지) into visit history records.

    Supports two formats:
    1. Single-visit document with labeled fields
    2. Multi-visit document with separator patterns (---, ===, numbered)
    """
    records = []
    warnings = []
    glossary_used = {}

    # Split into visit blocks (separated by ---, ===, or numbered patterns)
    blocks = _split_into_blocks(paragraphs)

    for block_idx, block in enumerate(blocks, start=1):
        fields = {}
        record_warnings = []
        raw_text = "\n".join(block)

        # Extract labeled fields
        for field_name, pattern in VISIT_FIELD_PATTERNS.items():
            for line in block:
                m = pattern.search(line)
                if m:
                    fields[field_name] = m.group(1).strip()
                    break

        # Extract date from the block
        visit_date = None
        if "일시" in fields:
            visit_date, _ = normalize_date(fields["일시"])
        if not visit_date:
            # Scan all lines for a date
            for line in block:
                d, conf = normalize_date(line)
                if d and conf > 0.5:
                    visit_date = d
                    break

        # Extract member reference
        member_ref = fields.get("방문대상")
        member_id_match = MEMBER_ID_RE.search(raw_text)
        member_id = member_id_match.group(0) if member_id_match else None

        # Normalize visitor role using glossary
        visitor = fields.get("방문자", "")
        for korean_term, english_term in glossary.items():
            if korean_term in visitor:
                glossary_used[korean_term] = english_term

        # Build the visit record (maps to members.yaml history entry)
        if not visit_date and not member_ref:
            if raw_text.strip():
                warnings.append(
                    f"Block {block_idx}: No date or member found — content: "
                    f"{raw_text[:100]}..."
                )
            continue

        confidence = 0.85
        if not visit_date:
            confidence -= 0.2
            record_warnings.append(f"Block {block_idx}: No date found")
        if not member_ref:
            confidence -= 0.15
            record_warnings.append(f"Block {block_idx}: No member reference found")

        record_fields = {
            "date": visit_date,
            "event": "pastoral_visitation",
            "member_name": member_ref,
            "member_id": member_id,
            "visitor": fields.get("방문자"),
            "location": fields.get("장소"),
            "reason": fields.get("사유"),
            "content": fields.get("내용"),
            "prayer_topics": fields.get("기도제목"),
            "follow_up": fields.get("후속조치"),
            "note": _build_visit_note(fields),
        }

        records.append({
            "fields": record_fields,
            "confidence": round(max(confidence, 0.3), 2),
            "source_block": block_idx,
            "notes": record_warnings,
        })
        warnings.extend(record_warnings)

    return records, warnings, glossary_used


def parse_meeting_minutes(paragraphs, glossary, source_file):
    """Parse meeting minutes (회의록) into structured records."""
    records = []
    warnings = []
    glossary_used = {}

    # Extract meeting metadata from first few paragraphs
    meeting_type = None
    meeting_date = None
    attendees = []
    agenda_items = []
    decisions = []

    current_section = None

    for line in paragraphs:
        # Detect meeting type
        for korean_term in ("당회", "제직회", "공동의회"):
            if korean_term in line:
                if korean_term in glossary:
                    meeting_type = glossary[korean_term]
                    glossary_used[korean_term] = meeting_type
                else:
                    meeting_type = korean_term
                break

        # Detect date
        if not meeting_date:
            d, conf = normalize_date(line)
            if d and conf > 0.5:
                meeting_date = d

        # Section detection
        if re.search(r"참석자|출석", line):
            current_section = "attendees"
            continue
        elif re.search(r"안건|의안|회의\s*안건", line):
            current_section = "agenda"
            continue
        elif re.search(r"결의|결정|의결", line):
            current_section = "decisions"
            continue

        # Collect section content
        if current_section == "attendees" and line.strip():
            attendees.append(line.strip())
        elif current_section == "agenda" and line.strip():
            agenda_items.append(line.strip())
        elif current_section == "decisions" and line.strip():
            decisions.append(line.strip())

    confidence = 0.8
    if not meeting_date:
        confidence -= 0.2
        warnings.append("No meeting date found in document")
    if not meeting_type:
        confidence -= 0.1
        warnings.append("Meeting type not identified")

    record_fields = {
        "date": meeting_date,
        "type": meeting_type or "unknown",
        "attendees": attendees,
        "agenda": agenda_items,
        "decisions": decisions,
        "full_text": "\n".join(paragraphs),
    }

    records.append({
        "fields": record_fields,
        "confidence": round(max(confidence, 0.3), 2),
        "source_block": 1,
        "notes": [],
    })

    return records, warnings, glossary_used


def parse_generic_document(paragraphs, glossary, source_file):
    """Fallback parser for unrecognized .docx/.pdf documents.

    Extracts any structured fields it can find and packages the full text.
    """
    records = []
    warnings = []
    glossary_used = {}

    # Extract any dates, phone numbers, member IDs from text
    full_text = "\n".join(paragraphs)
    dates = DATE_RE.findall(full_text)
    dates_ko = DATE_KO_RE.findall(full_text)
    phones = PHONE_RE.findall(full_text)
    member_ids = MEMBER_ID_RE.findall(full_text)

    # Look for glossary terms
    for korean_term, english_term in glossary.items():
        if korean_term in full_text:
            glossary_used[korean_term] = english_term

    record_fields = {
        "full_text": full_text,
        "extracted_dates": dates + [
            f"{y}-{int(m):02d}-{int(d):02d}" for y, m, d in dates_ko
        ],
        "extracted_phones": phones,
        "extracted_member_ids": member_ids,
        "glossary_terms_found": list(glossary_used.keys()),
        "paragraph_count": len(paragraphs),
    }

    confidence = 0.4  # Generic parse = low confidence
    warnings.append(
        f"Unrecognized document format. Extracted raw data from {len(paragraphs)} paragraphs. "
        "Manual review recommended."
    )

    records.append({
        "fields": record_fields,
        "confidence": round(confidence, 2),
        "source_block": 1,
        "notes": ["Generic parse — manual review required"],
    })

    return records, warnings, glossary_used


# ---------------------------------------------------------------------------
# Helper functions
# ---------------------------------------------------------------------------
def _split_into_blocks(paragraphs):
    """Split paragraphs into logical blocks (for multi-record documents)."""
    blocks = []
    current_block = []

    for line in paragraphs:
        # Separator patterns
        if re.match(r"^[-=]{3,}$", line) or re.match(r"^\d+\.\s*$", line):
            if current_block:
                blocks.append(current_block)
                current_block = []
            continue
        # Numbered visit entries: "1. 심방일지", "제1회 심방"
        if re.match(r"^(?:제?\d+[.회]\s*)", line) and current_block:
            blocks.append(current_block)
            current_block = []
        current_block.append(line)

    if current_block:
        blocks.append(current_block)

    # If no separators found, treat entire document as one block
    if not blocks:
        blocks = [paragraphs]

    return blocks


def _build_visit_note(fields):
    """Build a concise note string from extracted visitation fields."""
    parts = []
    if fields.get("사유"):
        parts.append(f"사유: {fields['사유']}")
    if fields.get("내용"):
        content = fields["내용"]
        if len(content) > 200:
            content = content[:200] + "..."
        parts.append(content)
    if fields.get("기도제목"):
        parts.append(f"기도제목: {fields['기도제목']}")
    return " | ".join(parts) if parts else None


# ---------------------------------------------------------------------------
# Route detection
# ---------------------------------------------------------------------------
def detect_file_type(filepath):
    """Detect which schema parser to use based on filename patterns."""
    basename = os.path.splitext(os.path.basename(filepath))[0]
    for pattern, info in FILE_ROUTING.items():
        if pattern in basename:
            return pattern, info
    return None, None


# ---------------------------------------------------------------------------
# Main parse entry point
# ---------------------------------------------------------------------------
def parse_file(filepath, glossary_path, staging_dir):
    """Parse a single semi-structured file and produce a staging JSON.

    Args:
        filepath: Path to .docx or .pdf file.
        glossary_path: Path to church-glossary.yaml.
        staging_dir: Directory for staging output.

    Returns:
        dict with parse results, or error dict.
    """
    if not os.path.isfile(filepath):
        return {"error": f"File not found: {filepath}"}

    ext = os.path.splitext(filepath)[1].lower()
    glossary = load_glossary(glossary_path)

    # Extract text
    try:
        if ext == ".docx":
            paragraphs = extract_text_docx(filepath)
        elif ext == ".pdf":
            paragraphs = extract_text_pdf(filepath)
        else:
            return {"error": f"Unsupported extension for Tier B: {ext}"}
    except RuntimeError as e:
        return {"error": str(e)}
    except Exception as e:
        return {"error": f"Failed to extract text from {filepath}: {e}"}

    if not paragraphs:
        return {"error": f"No text content extracted from {filepath}"}

    # Detect file type and route
    route_key, route_info = detect_file_type(filepath)

    if route_info:
        parser_name = route_info["parser"]
        target_file = route_info["target"]
        target_section = route_info["section"]
    else:
        parser_name = "parse_generic_document"
        target_file = "unknown"
        target_section = "unknown"

    # Execute parser
    parser_fn = {
        "parse_visitation": parse_visitation,
        "parse_meeting_minutes": parse_meeting_minutes,
        "parse_generic_document": parse_generic_document,
    }.get(parser_name, parse_generic_document)

    records, parse_warnings, glossary_used = parser_fn(
        paragraphs, glossary, filepath
    )

    # Build staging output
    confidences = [r["confidence"] for r in records]
    avg_confidence = sum(confidences) / len(confidences) if confidences else 0.0

    staging_data = {
        "source_file": filepath,
        "parser_tier": "B",
        "target_data_file": target_file,
        "target_section": target_section,
        "records": records,
        "glossary_mappings": glossary_used,
        "parse_warnings": parse_warnings,
        "text_stats": {
            "total_paragraphs": len(paragraphs),
            "total_characters": sum(len(p) for p in paragraphs),
        },
        "timestamp": datetime.now().isoformat(timespec="seconds"),
        "total_records": len(records),
        "average_confidence": round(avg_confidence, 2),
    }

    # Write staging file
    os.makedirs(staging_dir, exist_ok=True)
    basename = os.path.splitext(os.path.basename(filepath))[0]
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    staging_filename = f"tier_b_{basename}_{ts}.json"
    staging_path = os.path.join(staging_dir, staging_filename)

    with open(staging_path, "w", encoding="utf-8") as f:
        json.dump(staging_data, f, indent=2, ensure_ascii=False)

    staging_data["staging_file"] = staging_path
    return staging_data


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------
def main():
    parser = argparse.ArgumentParser(
        description="Tier B Parser — Parse semi-structured data (Word/PDF) for church admin"
    )
    parser.add_argument(
        "--file", required=True,
        help="Path to .docx or .pdf file to parse"
    )
    parser.add_argument(
        "--glossary", default="./data/church-glossary.yaml",
        help="Path to church-glossary.yaml (default: ./data/church-glossary.yaml)"
    )
    parser.add_argument(
        "--staging-dir", default="./inbox/staging",
        help="Directory for staging output (default: ./inbox/staging)"
    )
    parser.add_argument(
        "--json", action="store_true",
        help="Output full JSON result to stdout"
    )
    args = parser.parse_args()

    result = parse_file(args.file, args.glossary, args.staging_dir)

    if "error" in result:
        print(f"ERROR: {result['error']}", file=sys.stderr)
        if args.json:
            print(json.dumps(result, indent=2, ensure_ascii=False))
        sys.exit(1)

    print(f"Tier B parse complete: {result['total_records']} records extracted")
    print(f"  Source:     {result['source_file']}")
    print(f"  Target:     {result['target_data_file']} [{result['target_section']}]")
    print(f"  Confidence: {result['average_confidence']:.0%}")
    print(f"  Text:       {result['text_stats']['total_paragraphs']} paragraphs, "
          f"{result['text_stats']['total_characters']} chars")
    print(f"  Warnings:   {len(result['parse_warnings'])}")
    if result.get("staging_file"):
        print(f"  Staging:    {result['staging_file']}")

    if args.json:
        print(json.dumps(result, indent=2, ensure_ascii=False))


if __name__ == "__main__":
    main()
